import streamlit as st
import joblib
import os
import io
import pandas as pd
import re
import seaborn as sns
import matplotlib.pyplot as plt
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.svm import SVC
from sklearn.metrics import accuracy_score, classification_report
from sklearn.model_selection import train_test_split
from imblearn.over_sampling import RandomOverSampler
from sklearn.preprocessing import LabelEncoder
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS
import nltk
import spacy
import pdfplumber
import docx
from io import StringIO
from wordcloud import WordCloud
import numpy as np
import plotly.graph_objects as go
from bs4 import BeautifulSoup
import requests
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from collections import Counter
from datetime import datetime
from sklearn.neighbors import KNeighborsClassifier
from sklearn.ensemble import GradientBoostingClassifier
from sklearn.naive_bayes import MultinomialNB
from imblearn.over_sampling import SMOTE
from imblearn.under_sampling import RandomUnderSampler
from imblearn.combine import SMOTEENN
import shutil
import gdown
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
import sklearn
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload, MediaFileUpload
from googleapiclient.errors import HttpError
import tempfile
import time


# --- Fonction d'authentification ---
def authenticate_user():
    st.sidebar.title("Login")
    username = st.sidebar.text_input("Nom d'utilisateur", "")
    password = st.sidebar.text_input("Mot de passe", type="password")

    # Ajouter un utilisateur et un mot de passe pour la démonstration (à sécuriser davantage dans une vraie application)
    correct_username = "admin"
    correct_password = "123"

    error_message_shown = False  # Flag pour vérifier si l'erreur doit être affichée

    if username and password:  # Vérifier si les deux champs sont remplis
        if username != correct_username or password != correct_password:
            error_message_shown = True  # Marquer l'erreur pour qu'elle soit affichée après la tentative

    # Si l'utilisateur essaie de se connecter avec de mauvaises informations
    if error_message_shown:
        st.sidebar.error("Nom d'utilisateur ou mot de passe incorrect.")

    # Vérification des informations de connexion
    if username == correct_username and password == correct_password:
        return True
    else:
        return False


# --- Script 1: Resume Classification Model Training ---
def resume_classification():
    nltk.download("stopwords")

    # Stopwords en français et en anglais
    french_stop_words = set(stopwords.words("french"))
    english_stop_words = ENGLISH_STOP_WORDS  # Stopwords en anglais
    all_stop_words = french_stop_words.union(english_stop_words)

    # Fonction avancée de nettoyage du texte
    def clean_text(text):
        if pd.isna(text):  # Gérer les valeurs NaN
            return ""
        text = text.lower()
        text = re.sub(r"[^a-zA-Z\s]", "", text)  # Supprimer les caractères spéciaux et chiffres
        text = re.sub(r"\s+", " ", text).strip()  # Supprimer les espaces supplémentaires
        text = " ".join([word for word in text.split() if word not in all_stop_words])  # Supprimer les stopwords
        return text

    # Fonction d'authentification avec gestion automatique du refresh token
    def authenticate_google_drive():
        creds = None
        SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/drive.readonly']

        # Charger le token existant s'il existe
        if os.path.exists('token.json'):
            creds = Credentials.from_authorized_user_file('token.json', SCOPES)

        # Vérifier si le token est expiré et peut être rafraîchi
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())

        # Si aucune crédential valide n'est trouvée, demander une nouvelle authentification
        if not creds or not creds.valid:
            flow = InstalledAppFlow.from_client_secrets_file('credentialss.json', SCOPES)
            creds = flow.run_console(port=8080, access_type='offline', prompt='consent')

            # Sauvegarder le token pour réutilisation future
            with open('token.json', 'w') as token:
                token.write(creds.to_json())

        return build('drive', 'v3', credentials=creds)

    # Télécharger un fichier CSV à partir de Google Drive en mémoire
    def download_file_from_drive(file_id):
        service = authenticate_google_drive()

        # Utiliser l'API Google Drive pour télécharger le fichier
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()  # Flux en mémoire pour éviter le téléchargement local
        downloader = MediaIoBaseDownload(fh, request)

        done = False
        while done is False:
            status, done = downloader.next_chunk()
            print(f"Download {int(status.progress() * 100)}%.")
        
        fh.seek(0)  # Revenir au début du flux pour pouvoir lire
        return fh

    # URL du fichier CSV
    file_url = "https://drive.google.com/uc?id=1-5Hw-uq7-NFJjcU7LuD_pgK42yJc8lxR"
    file_id = file_url.split('id=')[1]

    # Télécharger le fichier CSV depuis Google Drive en mémoire
    file_stream = download_file_from_drive(file_id)

    # Lire le CSV directement depuis le flux en mémoire avec pandas
    df = pd.read_csv(file_stream, on_bad_lines='skip')


    # Suppression des lignes avec des valeurs manquantes
    df.dropna(subset=['category', 'text'], inplace=True)

    # Nettoyage avancé du texte
    df["clean_text"] = df["text"].apply(clean_text)

    # Encodage des catégories
    label_encoder = LabelEncoder()
    df["category_encoded"] = label_encoder.fit_transform(df["category"])

    # Transformation du texte en vecteurs TF-IDF (bigrammes + 5000 features)
    tfidf_vectorizer = TfidfVectorizer(max_features=5000, ngram_range=(1, 2))
    X = tfidf_vectorizer.fit_transform(df["clean_text"])
    y = df["category_encoded"]

    # Affichage de la distribution des classes avant l'équilibrage
    class_distribution_before = df['category'].value_counts().reset_index()
    class_distribution_before.columns = ['Category', 'Count']
    class_distribution_before['Status'] = 'Avant équilibrage'

    # Équilibrage des classes sous-représentées avec RandomOverSampler
    ros = RandomOverSampler(random_state=42)
    X_resampled, y_resampled = ros.fit_resample(X, y)

    # Affichage de la distribution des classes après l'équilibrage
    class_distribution_after = pd.Series(y_resampled).value_counts().reset_index()
    class_distribution_after.columns = ['Category', 'Count']
    class_distribution_after['Status'] = 'Après équilibrage'

    # Fusionner les deux distributions pour comparaison
    class_distribution_comparison = pd.concat([class_distribution_before, class_distribution_after], axis=0)

    # Interface Streamlit
    st.title("Entraînement des modèles de classification")


    # Séparation des données en train/test (80% train, 20% test)
    X_train, X_test, y_train, y_test = train_test_split(X_resampled, y_resampled, test_size=0.2, random_state=42)

    # Fonction d'évaluation des modèles
    def evaluate_model(name, model, X_test, y_test):
        y_pred = model.predict(X_test)
        accuracy = accuracy_score(y_test, y_pred)
        report = classification_report(y_test, y_pred, output_dict=True)
        
        precision = report["macro avg"]["precision"]
        recall = report["macro avg"]["recall"]
        f1 = report["macro avg"]["f1-score"]
        
        return accuracy, precision, recall, f1, report

    # Section pour les hyperparamètres
    st.sidebar.header("Paramètres des modèles")

    # Paramètres pour Logistic Regression
    logreg_max_iter = st.sidebar.number_input("Max Iterations (Logistic Regression)", min_value=100, max_value=5000, value=2000)
    logreg_class_weight = st.sidebar.selectbox("Poids des classes (Logistic Regression)", ["balanced", "None"])

    # Paramètres pour Random Forest
    rf_n_estimators = st.sidebar.selectbox("Nombre d'estimateurs (RandomForest)", [100, 200, 300])
    rf_max_depth = st.sidebar.selectbox("Profondeur maximale (RandomForest)", [10, 20, None])

    # Paramètres pour SVM
    svm_kernel = st.sidebar.selectbox("Noyau (SVM)", ["linear", "rbf"])
    svm_class_weight = st.sidebar.selectbox("Poids des classes (SVM)", ["balanced", "None"])

    # Paramètres pour KNeighbors
    knn_n_neighbors = st.sidebar.selectbox("Nombre de voisins (KNeighbors)", [3, 5, 7, 9])

    # Paramètres pour GradientBoosting
    gb_n_estimators = st.sidebar.selectbox("Nombre d'estimateurs (Gradient Boosting)", [100, 200, 300])
    gb_learning_rate = st.sidebar.number_input("Taux d'apprentissage (Gradient Boosting)", min_value=0.01, max_value=1.0, value=0.1)

    # Paramètres pour Naive Bayes
    nb_alpha = st.sidebar.number_input("Alpha (Naive Bayes)", min_value=0.1, max_value=2.0, value=1.0)

    # Variables pour stocker les résultats
    results = {
        'Model': [],
        'Accuracy': [],
        'Precision': [],
        'Recall': [],
        'F1-Score': []
    }

    best_model = None
    best_model_name = ""
    model_trained = False  # Flag pour suivre si les modèles ont été entraînés

    # Authentification et connexion à Google Drive
    drive_service = authenticate_google_drive()

    # ID du modèle dans Google Drive
    models_dir = "1OaKR_9g_gpLNI0pSYbNJKNdO4jn4d35z"  # L'ID du dossier Google Drive

    # Bouton pour entraîner les modèles
    if st.sidebar.button("Entraîner les modèles"):
        # Vérifier le répertoire Models et le créer s'il n'existe pas
        if not os.path.exists(models_dir):
            os.makedirs(models_dir)
            st.write(f"Répertoire '{models_dir}' créé.")
        else:
            st.write(f"Répertoire '{models_dir}' déjà existant.")
        with st.spinner("Les modèles sont en cours d'entraînement..."):
            # Régression logistique
            logreg_model = LogisticRegression(max_iter=logreg_max_iter, class_weight=logreg_class_weight)
            logreg_model.fit(X_train, y_train)
            acc_logreg, precision_logreg, recall_logreg, f1_logreg, report_logreg = evaluate_model("Logistic Regression", logreg_model, X_test, y_test)
            results['Model'].append('Logistic Regression')
            results['Accuracy'].append(acc_logreg)
            results['Precision'].append(precision_logreg)
            results['Recall'].append(recall_logreg)
            results['F1-Score'].append(f1_logreg)

            # Random Forest
            rf_model = RandomForestClassifier(n_estimators=rf_n_estimators, max_depth=rf_max_depth, random_state=42)
            rf_model.fit(X_train, y_train)
            acc_rf, precision_rf, recall_rf, f1_rf, report_rf = evaluate_model("Random Forest", rf_model, X_test, y_test)
            results['Model'].append('Random Forest')
            results['Accuracy'].append(acc_rf)
            results['Precision'].append(precision_rf)
            results['Recall'].append(recall_rf)
            results['F1-Score'].append(f1_rf)

            # SVM
            svm_model = SVC(kernel=svm_kernel, class_weight=svm_class_weight)
            svm_model.fit(X_train, y_train)
            acc_svm, precision_svm, recall_svm, f1_svm, report_svm = evaluate_model("SVM", svm_model, X_test, y_test)
            results['Model'].append('SVM')
            results['Accuracy'].append(acc_svm)
            results['Precision'].append(precision_svm)
            results['Recall'].append(recall_svm)
            results['F1-Score'].append(f1_svm)

            # KNeighbors
            knn_model = KNeighborsClassifier(n_neighbors=knn_n_neighbors)
            knn_model.fit(X_train, y_train)
            acc_knn, precision_knn, recall_knn, f1_knn, report_knn = evaluate_model("KNeighbors", knn_model, X_test, y_test)
            results['Model'].append('KNeighbors')
            results['Accuracy'].append(acc_knn)
            results['Precision'].append(precision_knn)
            results['Recall'].append(recall_knn)
            results['F1-Score'].append(f1_knn)

            # Gradient Boosting
            gb_model = GradientBoostingClassifier(n_estimators=gb_n_estimators, learning_rate=gb_learning_rate)
            gb_model.fit(X_train, y_train)
            acc_gb, precision_gb, recall_gb, f1_gb, report_gb = evaluate_model("Gradient Boosting", gb_model, X_test, y_test)
            results['Model'].append('Gradient Boosting')
            results['Accuracy'].append(acc_gb)
            results['Precision'].append(precision_gb)
            results['Recall'].append(recall_gb)
            results['F1-Score'].append(f1_gb)

            # Naive Bayes
            nb_model = MultinomialNB(alpha=nb_alpha)
            nb_model.fit(X_train, y_train)
            acc_nb, precision_nb, recall_nb, f1_nb, report_nb = evaluate_model("Naive Bayes", nb_model, X_test, y_test)
            results['Model'].append('Naive Bayes')
            results['Accuracy'].append(acc_nb)
            results['Precision'].append(precision_nb)
            results['Recall'].append(recall_nb)
            results['F1-Score'].append(f1_nb)

            # Création du DataFrame des résultats
            results_df = pd.DataFrame(results)

            # Affichage des résultats sous forme de tableau
            st.subheader("Tableau des Métriques des Modèles")
            st.dataframe(results_df)

            # Déterminer le meilleur modèle
            best_model_info = max([(logreg_model, "Logistic Regression", acc_logreg),
                                  (rf_model, "Random Forest", acc_rf),
                                  (svm_model, "SVM", acc_svm),
                                  (knn_model, "KNeighbors", acc_knn),
                                  (gb_model, "Gradient Boosting", acc_gb),
                                  (nb_model, "Naive Bayes", acc_nb)], key=lambda x: x[2])

            best_model_name = best_model_info[1]
            best_model = best_model_info[0]
            best_model_acc = best_model_info[2]

            st.write(f"**Meilleur modèle : {best_model_name} avec une précision de {best_model_acc:.4f}**")

            # Upload des modèles directement sur Google Drive
            try:
                st.write("🔄 Upload des modèles vers Google Drive...")

                # Fonction pour uploader un fichier sur Google Drive et supprimer l'ancien modèle s'il existe
                def upload_to_drive(model, drive_folder_id, filename):
                    # Vérification si un fichier avec le même nom existe déjà
                    query = f"'{drive_folder_id}' in parents and name = '{filename}'"
                    results = drive_service.files().list(q=query, fields="files(id, name)").execute()
                    files = results.get('files', [])

                    # Si le fichier existe, on le supprime
                    if files:
                        file_id = files[0]['id']
                        drive_service.files().delete(fileId=file_id).execute()
                        st.write(f"❌ Fichier {filename} déjà existant, il a été supprimé.")

                    # Sérialiser l'objet modèle en mémoire avec io.BytesIO
                    model_data = io.BytesIO()
                    joblib.dump(model, model_data)
                    model_data.seek(0)  # Revenir au début du fichier binaire

                    # Création du fichier pour Google Drive
                    file_metadata = {'name': filename, 'parents': [drive_folder_id]}
                    media = MediaIoBaseUpload(model_data, mimetype='application/octet-stream')
                    file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
                    return file['id']

                # Upload des modèles vers Google Drive
                model_id = upload_to_drive(best_model, models_dir, "best_model.pkl")
                tfidf_id = upload_to_drive(tfidf_vectorizer, models_dir, "tfidf_vectorizer.pkl")
                label_encoder_id = upload_to_drive(label_encoder, models_dir, "label_encoder.pkl")

                st.success("✅ Modèles sauvegardés sur Google Drive avec succès !")

                # Affichage des liens de téléchargement
                st.write(f"🔗 [Best Model](https://drive.google.com/file/d/{model_id}/view)")
                st.write(f"🔗 [TF-IDF Vectorizer](https://drive.google.com/file/d/{tfidf_id}/view)")
                st.write(f"🔗 [Label Encoder](https://drive.google.com/file/d/{label_encoder_id}/view)")

            except FileNotFoundError as fnf_error:
                st.error(f"❌ Erreur FileNotFoundError : {str(fnf_error)}")
            except PermissionError as perm_error:
                st.error(f"❌ Erreur PermissionError : {str(perm_error)}")
            except Exception as e:
                st.error(f"❌ Erreur inconnue lors de la sauvegarde du modèle : {str(e)}")


# --- Script 2: Automated CV Analysis and Job Match ---
def automated_cv_analysis():

    # Télécharger les stopwords
    nltk.download("stopwords")

    # Fonction pour extraire le texte d'un fichier (PDF, DOCX, texte, ou fichier Google Drive)
    def extract_text(uploaded_file=None, file_id=None, service=None):
        if uploaded_file:  # Si un fichier est téléchargé via Streamlit
            file_name = uploaded_file.name.lower()

            if file_name.endswith(".pdf"):
                with pdfplumber.open(uploaded_file) as pdf:
                    return " ".join([page.extract_text() for page in pdf.pages if page.extract_text()])

            elif file_name.endswith(".docx"):
                doc = docx.Document(uploaded_file)
                return " ".join([para.text for para in doc.paragraphs])

            else:  # Pour les fichiers texte ou autres formats
                try:
                    return uploaded_file.read().decode("utf-8")
                except UnicodeDecodeError:
                    return "Erreur de décodage, le fichier n'est pas en UTF-8."
        elif file_id and service:  # Si un fichier est sur Google Drive
            request = service.files().get_media(fileId=file_id)
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()

            file_io.seek(0)
            
            # On détermine le type du fichier en fonction de son extension
            file_info = service.files().get(fileId=file_id).execute()
            mime_type = file_info['mimeType']

            if mime_type == 'application/pdf':
                with pdfplumber.open(file_io) as pdf:
                    return " ".join([page.extract_text() for page in pdf.pages if page.extract_text()])

            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                doc = docx.Document(file_io)
                return " ".join([para.text for para in doc.paragraphs])

            elif mime_type == 'text/plain':
                return file_io.read().decode('utf-8')

            else:
                return "Format de fichier non pris en charge pour l'extraction de texte."

        return "Aucun fichier à traiter."

    # 📌 Fonction de nettoyage du texte
    def clean_text(text):
        text = text.lower()
        text = re.sub(r"\d+", "", text)  # Supprimer les chiffres
        text = re.sub(r"[^\w\s]", "", text)  # Supprimer la ponctuation
        text = " ".join([word for word in text.split() if word not in stopwords.words("french")])
        return text

    # 📌 Fonction pour extraire les technologies du CV sans doublons et leur fréquence
    def extract_technologies_with_count(text):
        # Liste complète des compétences à extraire
        tech_keywords = [
            # Data Engineer
            "Python", "SQL", "Spark", "AWS", "Kafka", "Airflow", "Snowflake", "Redshift", "Databricks", "Docker", "Kubernetes", "Jenkins", "ETL", "Pipeline",

            # Data Scientist
            "Machine Learning", "Deep Learning", "NLP", "Computer Vision", "TensorFlow", "PyTorch", "Keras", "Scikit-learn", "Transformers", "BERT", "LSTM", "GANs", "Reinforcement Learning",

            # Data Analyst
            "Tableau", "Power BI", "SQL", "Excel", "Looker", "Google Data Studio", "DAX", "Pandas", "Matplotlib", "Reporting", "ETL", "Dashboard", "Visualization",

            # DevOps
            "Docker", "Kubernetes", "Terraform", "CI/CD", "Ansible", "Jenkins", "Git", "Helm", "Prometheus", "Grafana", "ArgoCD", "Istio", "OpenShift", "Infrastructure as Code",

            # Backend Developer
            "Java", "Spring Boot", "Microservices", "Hibernate", "REST API", "JPA", "SQL", "NoSQL", "RabbitMQ", "Kafka", "GraphQL", "WebFlux", "Docker", "Kubernetes",

            # Frontend Developer
            "React", "Angular", "Vue.js", "JavaScript", "TypeScript", "HTML", "CSS", "Redux", "Next.js", "Nuxt.js", "Tailwind", "Material-UI", "Cypress", "Jest", "Storybook", "Webpack",

            # Mobile Developer
            "Swift", "Kotlin", "Flutter", "React Native", "Android", "iOS", "Jetpack Compose", "SwiftUI", "Objective-C", "Dart", "Mobile UI/UX", "GraphQL", "Firebase",

            # Big Data Engineer
            "Hadoop", "Hive", "Pig", "Scala", "Spark", "Presto", "Flink", "HBase", "Cassandra", "ElasticSearch", "MapReduce", "Delta Lake", "Kudu", "YARN", "Zookeeper",

            # Cybersecurity Analyst
            "Cybersecurity", "Penetration Testing", "SIEM", "IDS/IPS", "Firewall", "Ethical Hacking", "Kali Linux", "Metasploit", "OWASP", "Burp Suite", "Nmap", "Security Compliance",

            # ETL Developer
            "ETL", "Talend", "SSIS", "Informatica", "DataStage", "Azure Data Factory", "Apache Nifi", "Pentaho", "Snowflake", "SQL", "Data Warehouse", "OLAP",

            # Database Administrator
            "Oracle", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Cassandra", "MariaDB", "CockroachDB", "Elasticsearch", "TimescaleDB", "SQL", "NoSQL", "Database Replication", "Indexing",

            # CRM Consultant
            "Salesforce", "SAP", "ERP", "CRM", "Dynamics 365", "Zoho CRM", "HubSpot", "Workday", "NetSuite", "ServiceNow", "Business Process Automation", "RPA",

            # Business Intelligence Analyst
            "Excel", "VBA", "R", "Power BI", "Tableau", "SAS", "QlikView", "SQL", "Alteryx", "ETL", "Reporting", "KPI", "Business Intelligence", "Dashboarding",

            # Cloud Engineer
            "GCP", "AWS", "Azure", "Terraform", "Ansible", "CloudFormation", "Kubernetes", "Lambda", "Serverless", "DevOps", "Cloud Security", "IAM", "Networking",

            # AI Researcher
            "Computer Vision", "OpenCV", "YOLO", "Deep Learning", "TensorFlow", "PyTorch", "GANs", "Image Segmentation", "Object Detection", "Image Processing",

            # Blockchain Developer
            "Blockchain", "Ethereum", "Smart Contracts", "Solidity", "Hyperledger", "Polkadot", "Binance Smart Chain", "DeFi", "dApps", "NFTs", "Consensus Mechanisms",

            # IoT Engineer
            "IoT", "MQTT", "Edge Computing", "Raspberry Pi", "Arduino", "LoRaWAN", "Zigbee", "Smart Devices", "Industrial IoT", "IoT Security",

            # Network Engineer
            "Networking", "TCP/IP", "BGP", "OSPF", "SDN", "Cisco", "Juniper", "Wireshark", "Routing", "Switching", "Network Security", "VLAN", "MPLS",

            # IT Project Manager
            "Project Management", "Agile", "Scrum", "PMP", "Kanban", "Jira", "Confluence", "SAFe", "Prince2", "Risk Management", "Stakeholder Management", "Product Ownership"
        ]
        
        # Créer une expression régulière pour chaque technologie (en échappant les espaces et les parenthèses)
        tech_regex = [re.escape(tech.lower()) for tech in tech_keywords]
        combined_regex = r'\b(' + r'|'.join(tech_regex) + r')\b'

        # Trouver toutes les correspondances dans le texte (insensible à la casse)
        found_technologies = re.findall(combined_regex, text.lower())

        # Compter la fréquence des technologies, sans doublons
        tech_counts = Counter(found_technologies)

        # Retourner la liste des technologies avec leur fréquence sous forme de dictionnaire
        return dict(tech_counts)

    # 📌 Fonction pour détecter le niveau d'expérience

    def detect_experience(cv_text):
        cv_text = cv_text.lower()

        # Détecter les années d'expérience explicites
        match = re.search(r'(\d+)\s*(ans|years)', cv_text)
        years_of_experience = int(match.group(1)) if match else 0

        # Listes de mots-clés pour classifier
        junior_keywords = ["débutant", "junior", "stage", "apprentissage", "assistant"]
        senior_keywords = ["confirmé", "middle", "expérience significative", "3 ans", "4 ans", "5 ans", "6 ans", "7 ans"]
        expert_keywords = ["lead", "manager", "expert", "architecte", "10 ans", "15 ans", "principal"]

        # Déterminer le niveau basé sur les années détectées
        if years_of_experience >= 8:
            return "Expert"
        elif 3 <= years_of_experience < 8:
            return "Senior"
        elif years_of_experience > 0:
            return "Junior"

        # Si aucune année n'est détectée, on vérifie les mots-clés supplémentaires
        if any(word in cv_text for word in expert_keywords):
            return "Expert"
        elif any(word in cv_text for word in senior_keywords):
            return "Senior"
        elif any(word in cv_text for word in junior_keywords):
            return "Junior"

        # Par défaut, classer comme Junior si aucune correspondance n'est trouvée
        return "Junior"

    # 📌 Fonction pour créer un nuage de mots
    def generate_wordcloud(text):
        wordcloud = WordCloud(width=800, height=400, background_color="white").generate(text)
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation="bilinear")
        plt.axis("off")
        st.pyplot(plt)

    # 📌 Fonction pour générer un graphique radar des compétences
    def plot_radar_chart(user_skills, ideal_profile):
        labels = list(set(user_skills.keys()).union(set(ideal_profile.keys())))
        user_values = [user_skills.get(skill, 0) for skill in labels]
        ideal_values = [ideal_profile.get(skill, 0) for skill in labels]
        
        fig = go.Figure()
        fig.add_trace(go.Scatterpolar(r=user_values, theta=labels, fill='toself', name='CV'))
        fig.add_trace(go.Scatterpolar(r=ideal_values, theta=labels, fill='toself', name='Profil Idéal'))
        
        fig.update_layout(polar=dict(radialaxis=dict(visible=True)), showlegend=True)
        st.plotly_chart(fig)

    # Fonction d'authentification avec gestion automatique du refresh token
    def authenticate_google_drive():
        creds = None
        SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/drive.readonly']

        # Charger le token existant s'il existe
        if os.path.exists('token.json'):
            creds = Credentials.from_authorized_user_file('token.json', SCOPES)

        # Vérifier si le token est expiré et peut être rafraîchi
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())

        # Si aucune crédential valide n'est trouvée, demander une nouvelle authentification
        if not creds or not creds.valid:
            flow = InstalledAppFlow.from_client_secrets_file('credentialss.json', SCOPES)
            creds = flow.run_console(port=8080, access_type='offline', prompt='consent')

            # Sauvegarder le token pour réutilisation future
            with open('token.json', 'w') as token:
                token.write(creds.to_json())

        return build('drive', 'v3', credentials=creds)

    # Fonction pour chercher un fichier par son nom dans un répertoire spécifique sur Google Drive
    def find_file_in_drive(service, file_name, folder_id):
        query = f"'{folder_id}' in parents and name='{file_name}'"
        results = service.files().list(q=query).execute()
        files = results.get('files', [])
        if not files:
            return None
        return files[0]  # Retourner le premier fichier trouvé (si plusieurs fichiers ont le même nom, retourne le premier)

    # Fonction pour télécharger un fichier vers un autre répertoire de Google Drive
    def upload_file_to_drive(service, file_name, file_content, folder_id):
        # Créer un répertoire temporaire dédié
        temp_dir = os.path.join(os.getcwd(), 'temp_uploads')
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)  # Créer le répertoire s'il n'existe pas

        temp_file_path = os.path.join(temp_dir, file_name)
        
        # Créer un fichier temporaire avec delete=False
        with open(temp_file_path, 'wb') as temp_file:
            temp_file.write(file_content.read())  # Écrire le contenu du fichier téléchargé
            temp_file.flush()  # Assurez-vous que tout est écrit dans le fichier

        # Créer le média pour l'upload
        media = MediaFileUpload(temp_file_path, mimetype='application/octet-stream')

        # Metadata du fichier à télécharger
        file_metadata = {
            'name': file_name,
            'parents': [folder_id]
        }

        # Télécharger le fichier vers Google Drive
        uploaded_file = service.files().create(
            media_body=media,
            body=file_metadata,
            fields='id'
        ).execute()

        # Attendre un petit moment avant de supprimer le fichier pour éviter des conflits
        time.sleep(1)

        # Essayer plusieurs fois de supprimer le fichier (si le fichier est encore utilisé)
        for _ in range(3):
            try:
                os.remove(temp_file_path)  # Supprimer le fichier temporaire
                break  # Sortir de la boucle dès que le fichier est supprimé avec succès
            except PermissionError:
                time.sleep(1)  # Attendre un peu avant de réessayer

        return uploaded_file


    # Fonction pour télécharger un fichier depuis Google Drive
    def download_file_from_drive(service, file_id):
        request = service.files().get_media(fileId=file_id)
        file = io.BytesIO()
        downloader = MediaIoBaseDownload(file, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        file.seek(0)  # Remettre le curseur au début du fichier
        return file
    
    # Fonction pour vérifier si un fichier existe déjà dans le dossier Output sur Google Drive
    def file_exists_in_output(service, file_name, output_folder_id):
        query = f"'{output_folder_id}' in parents and name='{file_name}'"
        results = service.files().list(q=query).execute()
        files = results.get('files', [])
        return len(files) > 0  # Si la longueur de la liste est > 0, cela signifie qu'un fichier existe déjà

    # Fonction pour lister les fichiers dans un dossier spécifique
    def list_files_in_folder(service, folder_id):
        results = service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields="files(id, name, mimeType)"
        ).execute()
        return results.get('files', [])
    
    # Fonction pour vérifier si un fichier existe déjà dans le dossier Google Drive
    def file_exists_in_folder(service, folder_id, file_name):
        files_in_folder = list_files_in_folder(service, folder_id)
        for file in files_in_folder:
            if file['name'] == file_name:
                return True
        return False
    
    # Fonction pour vérifier si un fichier existe dans Google Drive avant de le télécharger
    def file_exists(service, file_id):
        try:
            # Tentative de récupérer les informations du fichier
            file_info = service.files().get(fileId=file_id).execute()
            return True  # Si l'ID est valide, le fichier existe
        except HttpError as error:
            if error.resp.status == 404:
                return False  # Si le fichier n'existe pas
            else:
                raise error  # Si une autre erreur survient, la lever

    # Fonction pour télécharger un fichier depuis Google Drive
    def download_file(service, file_id):
        if file_exists(service, file_id):
            request = service.files().get_media(fileId=file_id)
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()

            file_io.seek(0)  # Revenir au début du fichier
            return file_io  # Retourner le fichier téléchargé sous forme de BytesIO
        else:
            st.error(f"Le fichier avec l'ID {file_id} n'a pas été trouvé sur Google Drive.")
            return None
       
    # Fonction pour comparer les technologies du CV avec celles de l'offre d'emploi
    def compare_technologies_with_cvs(job_description):
        job_tech_counts = extract_technologies_with_count(job_description)
        matching_cvs = []

        file_url = "https://drive.google.com/uc?id=1-5Hw-uq7-NFJjcU7LuD_pgK42yJc8lxR"
        df = pd.read_csv(gdown.download(file_url, quiet=True), on_bad_lines='skip')

        cvs_folder_id = "1HNftka5J3QEf2mV2ZZsRMQMNHgpPD_K0"
        service = authenticate_google_drive()

        if service:
            for _, row in df.iterrows():
                cv_file_name = row['cv_filename']
                cv_text = row['text']
                tech_counts = extract_technologies_with_count(cv_text)

                common_tech = set(tech_counts.keys()).intersection(set(job_tech_counts.keys()))
                if common_tech:
                    similarity_score = compare_texts(cv_text, job_description)
                    matching_cvs.append({
                        'cv_file': cv_file_name,
                        'common_tech': common_tech,
                        'tech_counts': tech_counts,
                        'job_tech_counts': job_tech_counts,
                        'similarity_score': similarity_score,
                    })

        return matching_cvs
    
    # Fonction pour exporter les CVs correspondants dans un dossier Google Drive
    def export_matching_cvs(matching_cvs, output_folder_id):
        service = authenticate_google_drive()
        if service:
            for match in matching_cvs[:5]:
                cv_file_name = match['cv_file']
                file = find_file_in_drive(service, cv_file_name, "1HNftka5J3QEf2mV2ZZsRMQMNHgpPD_K0")

                if file:
                    file_content = download_file_from_drive(service, file['id'])
                    if not file_exists_in_output(service, cv_file_name, output_folder_id):
                        uploaded_file = upload_file_to_drive(service, cv_file_name, file_content, output_folder_id)
                        download_link = f"https://drive.google.com/uc?id={uploaded_file['id']}"
                        st.success(f"Le fichier '{cv_file_name}' a été exporté vers Google Drive.")
                        st.markdown(f"[Télécharger le fichier ici]({download_link})")
                    else:
                        st.warning(f"Le fichier '{cv_file_name}' existe déjà dans le dossier de sortie.")
                        file = find_file_in_drive(service, cv_file_name, output_folder_id)
                        download_link = f"https://drive.google.com/uc?id={file['id']}"
                        st.markdown(f"[Télécharger le fichier ici]({download_link})")
                else:
                    st.error(f"Le fichier '{cv_file_name}' n'a pas été trouvé dans Google Drive.")

    # 📌 Fonction pour comparer les technologies détectées avec l'offre d'emploi
    def compare_technologies_with_job(tech_counts, job_description):
        # Extraire les technologies de l'offre d'emploi
        job_tech_counts = extract_technologies_with_count(job_description)

        # Calculer le nombre de technologies communes entre le CV et l'offre d'emploi
        common_tech = set(tech_counts.keys()).intersection(set(job_tech_counts.keys()))
        common_tech_count = sum([min(tech_counts[tech], job_tech_counts[tech]) for tech in common_tech])

        return common_tech_count

    # 📌 Fonction pour obtenir les mots les plus fréquents dans un texte
    def get_top_keywords(text, num_keywords=5):
        words = text.split()
        words = [word for word in words if word not in stopwords.words("french")]  # Supprimer les stopwords
        word_counts = Counter(words)
        return word_counts.most_common(num_keywords)

    # 📌 Fonction pour créer un nuage de mots à partir des mots les plus fréquents
    def generate_top_keywords_wordcloud(text):
        # Extraire les 5 mots les plus fréquents
        top_keywords = get_top_keywords(text, num_keywords=5)
        
        # Créer un dictionnaire avec les mots et leurs fréquences
        word_freq = dict(top_keywords)

        # Générer le nuage de mots avec les fréquences
        wordcloud = WordCloud(width=800, height=400, background_color="white").generate_from_frequencies(word_freq)

        # Affichage du nuage de mots
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation="bilinear")
        plt.axis("off")
        st.pyplot(plt)

    # 📌 Fonction pour comparer les textes
    def compare_texts(text1, text2):
        # Vectorisation des textes
        vectorizer = CountVectorizer().fit_transform([text1, text2])
        cosine_sim = cosine_similarity(vectorizer[0:1], vectorizer[1:2])
        
        return cosine_sim[0][0]

    def download_from_drive(file_id, local_path):
        # Télécharger un fichier depuis Google Drive
        url = f'https://drive.google.com/uc?export=download&id={file_id}'
        response = requests.get(url)
        
        if response.status_code == 200:
            with open(local_path, 'wb') as f:
                f.write(response.content)
            return True
        else:
            return False

    def get_latest_file_ids(drive_service, folder_id):
        # Liste des fichiers dans le dossier spécifié, triés par date de création (descendant)
        results = drive_service.files().list(
            q=f"'{folder_id}' in parents",  # Rechercher les fichiers dans ce dossier
            fields="files(id, name, createdTime)",  # Récupérer l'ID, le nom et la date de création des fichiers
            orderBy="createdTime desc"  # Tri par date de création décroissante
        ).execute()

        files = results.get('files', [])

        if not files:
            print("Aucun fichier trouvé.")
            return None, None, None

        # Trier les fichiers par date de création (au cas où ce n'est pas déjà fait)
        files_sorted = sorted(files, key=lambda x: x['createdTime'], reverse=True)

        # Extraire les IDs des derniers fichiers
        model_id = next((file['id'] for file in files_sorted if 'best_model.pkl' in file['name']), None)
        tfidf_id = next((file['id'] for file in files_sorted if 'tfidf_vectorizer.pkl' in file['name']), None)
        label_encoder_id = next((file['id'] for file in files_sorted if 'label_encoder.pkl' in file['name']), None)

        return model_id, tfidf_id, label_encoder_id


    # Fonction pour télécharger un fichier depuis Google Drive
    def download_from_drive(drive_service, file_id, local_path):
        request = drive_service.files().get_media(fileId=file_id)
        fh = io.FileIO(local_path, 'wb')
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
            print(f"Téléchargement {int(status.progress() * 100)}%.")
        return local_path

    # Fonction de prédiction
    def predict_cv(cv_text):
        # 🔹 Authentification et connexion à Google Drive
        drive_service = authenticate_google_drive()

        # 🔹 ID du dossier Google Drive où les modèles sont sauvegardés
        models_dir = "1OaKR_9g_gpLNI0pSYbNJKNdO4jn4d35z"

        # 🔹 Récupérer les IDs des derniers modèles sauvegardés depuis Google Drive
        model_id, tfidf_id, label_encoder_id = get_latest_file_ids(drive_service, models_dir)

        # Vérification que tous les modèles existent
        if not all([model_id, tfidf_id, label_encoder_id]):
            st.error("❌ Certains fichiers modèles sont manquants dans Google Drive.")
            return

        # 🔹 Télécharger les modèles depuis Google Drive
        local_model_path = "best_model.pkl"
        local_tfidf_path = "tfidf_vectorizer.pkl"
        local_label_encoder_path = "label_encoder.pkl"

        download_from_drive(drive_service, model_id, local_model_path)
        download_from_drive(drive_service, tfidf_id, local_tfidf_path)
        download_from_drive(drive_service, label_encoder_id, local_label_encoder_path)

        # 🔹 Charger les modèles téléchargés
        model_category = joblib.load(local_model_path)
        vectorizer = joblib.load(local_tfidf_path)
        label_encoder = joblib.load(local_label_encoder_path)

        # Traitement du texte du CV
        cleaned_text = clean_text(cv_text)
        features = vectorizer.transform([cleaned_text])

        # Prédiction de la catégorie et de l'expérience
        predicted_category_encoded = model_category.predict(features)[0]
        predicted_category = label_encoder.inverse_transform([predicted_category_encoded])[0]
        predicted_experience = detect_experience(cv_text)

        return predicted_category, predicted_experience



    # 📌 Interface Streamlit
    st.title("📄🤖 CV Analysis & Job Match")
    st.subheader("Chargez un CV et obtenez la prédiction du métier et de l'expérience.")

    uploaded_file = st.file_uploader("📤 Importer un fichier (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

    # 🔥 L'utilisateur peut définir son propre profil idéal
    st.subheader("🔧 Définissez votre profil idéal")
    available_skills = ["Python","SQL","Sql","Spark","AWS","Kafka","Airflow","Snowflake","Redshift","Databricks","Docker","Kubernetes","Jenkins","ETL","Pipeline","Machine Learning",
            "Deep Learning","NLP","Computer Vision","TensorFlow","PyTorch","Keras","Scikit-learn","Transformers","BERT","LSTM","GANs","Reinforcement Learning","Tableau",
            "Power BI","Excel","Looker","Google Data Studio","DAX","Pandas","Matplotlib","Reporting","Dashboard","Visualization","Terraform","CI/CD","Ansible","Git","Helm","Prometheus","Grafana","ArgoCD","Istio",
            "OpenShift","Infrastructure as Code","Java","Spring Boot","Microservices","Hibernate","REST API","JPA","NoSQL","RabbitMQ","GraphQL","WebFlux","React","Angular","Vue.js","JavaScript","TypeScript","HTML","CSS","Redux","Next.js","Nuxt.js","Tailwind",
            "Material-UI","Cypress","Jest","Storybook","Webpack","Swift","Kotlin","Flutter","React Native","Android","iOS","Jetpack Compose","SwiftUI","Objective-C","Dart",
            "Mobile UI/UX","Firebase","Hadoop","Hive","Pig","Scala","Presto","Flink","HBase","Cassandra","ElasticSearch","MapReduce","Delta Lake","Kudu","YARN","Zookeeper","Cybersecurity","Penetration Testing","SIEM","IDS/IPS","Firewall","Ethical Hacking","Kali Linux","Metasploit","OWASP",
            "Burp Suite","Nmap","Security Compliance","Talend","SSIS","Informatica","DataStage","Azure Data Factory","Apache Nifi","Pentaho",
            "Data Warehouse","OLAP","Oracle","MySQL","PostgreSQL","MongoDB","Redis","MariaDB","CockroachDB","Elasticsearch","TimescaleDB",
            "Database Replication","Indexing","Salesforce","SAP","ERP","CRM","Dynamics 365","Zoho CRM","HubSpot","Workday","NetSuite","ServiceNow",
            "Business Process Automation","RPA","VBA","R","SAS","QlikView","Alteryx","KPI","Business Intelligence","Dashboarding","GCP","Azure","CloudFormation","Lambda","Serverless","DevOps","Cloud Security","IAM","Networking","OpenCV","YOLO",
            "Image Segmentation","Object Detection","Image Processing","Blockchain","Ethereum","Smart Contracts","Solidity","Hyperledger","Polkadot",
            "Binance Smart Chain","DeFi","dApps","NFTs","Consensus Mechanisms","IoT","MQTT","Edge Computing","Raspberry Pi","Arduino","LoRaWAN","Zigbee","Smart Devices",
            "Industrial IoT","IoT Security","TCP/IP","BGP","OSPF","SDN","Cisco","Juniper","Wireshark","Routing","Switching","Network Security","VLAN","MPLS",
            "Project Management","Agile","Scrum","PMP","Kanban","Jira","Confluence","SAFe","Prince2","Risk Management","Stakeholder Management","Product Ownership"]

    selected_skills = st.multiselect("Sélectionnez les compétences clés :", available_skills)
    ideal_profile = {skill: 1 for skill in selected_skills}

    # Exemple d'URL à afficher sous le champ de saisie (non cliquable)
    st.markdown("""
    *Exemple d'URL de la plateforme Emploi.ma :*  
    https://www.emploi.ma/offre-emploi-maroc/senior-data-engineers-hf-casablanca-rabat-8802045
    """)

    if uploaded_file:
        with st.spinner("🔍 Analyse en cours..."):
            cv_text = extract_text(uploaded_file)
            cleaned_cv_text = clean_text(cv_text)

            st.subheader("🔍 Contenu du CV :")
            st.text_area("Texte extrait :", cv_text, height=200)

            category, experience = predict_cv(cv_text)
            st.subheader(f"🎯 Métier : **{category}**")

            # Extraction des technologies et comptage
            tech_counts = extract_technologies_with_count(cv_text)

            if tech_counts:
                st.subheader("🛠️ Technologies détectées et leurs fréquences :")
                
                # Trier les technologies par fréquence (du plus élevé au plus bas)
                sorted_tech_counts = sorted(tech_counts.items(), key=lambda x: x[1], reverse=True)
                
                # Prendre uniquement les 10 premiers éléments
                top_10_tech = sorted_tech_counts[:10]

                # Appliquer la transformation pour mettre la première lettre en majuscule
                top_10_tech = [(tech.capitalize(), count) for tech, count in top_10_tech]
                
                # Affichage sous forme de tableau
                tech_data = [{"Technologie": tech, "Fréquence": count} for tech, count in top_10_tech]
                st.table(tech_data)
                
                # Affichage avec un graphique à barres
                tech_names = [tech for tech, count in top_10_tech]
                tech_frequencies = [count for tech, count in top_10_tech]
                
            else:
                st.subheader("🛠️ Aucune technologie détectée")

            # 🔥 Ajout du graphique radar
            st.subheader("📊 Visualisation des compétences")

            # Prendre uniquement les 10 technologies les plus fréquentes
            top_10_tech = sorted(tech_counts.items(), key=lambda x: x[1], reverse=True)[:10]

            # Appliquer la transformation pour mettre la première lettre en majuscule
            top_10_tech = [(tech.capitalize(), count) for tech, count in top_10_tech]
            
            user_skills = {tech: count for tech, count in top_10_tech}  # Utiliser seulement les 10 meilleures compétences

            # Visualisation avec le graphique radar
            plot_radar_chart(user_skills, ideal_profile)

            tech_counts = extract_technologies_with_count(cv_text)

            job_description = st.text_area("Entrez le texte de l'offre d'emploi")
            if job_description:
                st.subheader("🔍 Comparaison des technologies de l'offre avec les CVs")
                matching_cvs = compare_technologies_with_cvs(job_description)

                if matching_cvs:
                    matching_cvs_sorted = sorted(matching_cvs, key=lambda x: x['similarity_score'], reverse=True)[:5]
                    data = []
                    for match in matching_cvs_sorted:
                        data.append({
                            "Nom du CV": match["cv_file"],
                            "Technologies communes": ", ".join(match["common_tech"]),
                            "Score de Similarité": round(match["similarity_score"], 2)
                        })
                    st.write(pd.DataFrame(data).sort_values(by="Score de Similarité", ascending=False))

                    output_folder_id = '13mvWA3lyVWq1VwLAeDsNFSY0GslnRp2v'
                    export_matching_cvs(matching_cvs_sorted, output_folder_id)
                else:
                    st.write("Aucune technologie commune trouvée dans les CVs.")
            else:
                st.error("Veuillez entrer le texte de l'offre d'emploi.")

def set_sidebar_style():
    st.markdown("""
    <style>
        /* Style pour la barre latérale */
        .sidebar .sidebar-content {
            background-color: #000000; /* Arrière-plan noir pour la sidebar */
            color: white;
        }

        /* Style du titre dans la sidebar */
        .sidebar .sidebar-content h1, .sidebar .sidebar-content h2 {
            color: #ffffff;
            font-family: 'Arial', sans-serif;
        }

        /* Style des radio buttons */
        .stRadio>div {
            background-color: #333333; /* Fond gris foncé pour les boutons */
            color: white; 
            border: 2px solid #ffffff;  /* Bordure blanche */
            border-radius: 8px;
            padding: 12px;
            font-size: 18px;
            font-weight: bold;
            text-align: center;
            margin-bottom: 10px;
        }

        /* Style quand le bouton est survolé */
        .stRadio>div:hover {
            background-color: #555555; /* Changer la couleur de fond au survol */
        }

        /* Style du texte dans les boutons radio */
        .stRadio>div>label {
            color: white; /* Texte blanc pour les labels */
        }
    </style>
    """, unsafe_allow_html=True)

# --- Fonction principale avec les onglets stylisés ---
def main():
    set_sidebar_style()  # Appliquer les styles

    st.sidebar.title("Navigation")

    # Utiliser des boutons radio pour les onglets (afin de mieux les styliser)
    option = st.sidebar.radio(
        "Choisissez un onglet",
        ("Resume Classification", "CV Analysis and Job Match"),
        index=0,  # Index de l'onglet par défaut
        key="sidebar_radio"
    )

    # Sécuriser l'option "Resume Classification"
    if option == "Resume Classification":
        # Demander l'authentification avant d'accéder à cette option
        if authenticate_user():
            resume_classification()  # Si authentification réussie, afficher le contenu
        else:
            st.title("📂✨ Smart Resume Classification")
            st.write("Accès non autorisé. Veuillez entrer vos informations de connexion.")
    
    elif option == "CV Analysis and Job Match":
        automated_cv_analysis()  # Cette option reste publique

if __name__ == "__main__":
    main()
